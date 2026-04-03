"""
Attachment Processor Module

Handles downloading and processing of various file types from Discord attachments.
Supports text files, images, PDFs, DOCX, and other document types.
"""

import asyncio
import base64
import logging
from typing import Dict, List, Optional, Any, Tuple
from urllib.parse import urlparse
import io

import aiohttp
from utils.http_client import create_http_session

logger = logging.getLogger(__name__)


class AttachmentProcessor:
    """Processes various types of Discord attachments."""
    
    # Text file extensions that can be read directly
    TEXT_EXTENSIONS = {
        '.txt', '.md', '.json', '.xml', '.csv', '.yaml', '.yml',
        '.log', '.py', '.js', '.ts', '.jsx', '.tsx', '.html', '.css',
        '.scss', '.sass', '.java', '.c', '.cpp', '.h', '.hpp',
        '.go', '.rs', '.rb', '.php', '.sh', '.bash', '.sql',
        '.ini', '.conf', '.cfg', '.toml', '.env'
    }
    
    # Image extensions (handled by ImageProcessor)
    IMAGE_EXTENSIONS = {
        '.jpg', '.jpeg', '.png', '.gif', '.webp', '.bmp'
    }
    
    # Document extensions with text extraction
    DOCUMENT_EXTENSIONS = {
        '.pdf': 'pdf',
        '.docx': 'docx',
        '.doc': 'doc'
    }
    
    # Other document types (metadata only)
    OTHER_DOCUMENT_EXTENSIONS = {
        '.xlsx', '.xls', '.pptx', '.ppt', '.odt', '.ods', '.odp'
    }
    
    DISCORD_CDN_DOMAINS = [
        "cdn.discordapp.com",
        "media.discordapp.net"
    ]
    
    def __init__(self):
        self.download_timeout = 20  # seconds
        self.max_text_size_kb = 100
        self.max_image_size_mb = 20
        self.max_document_size_mb = 50
    
    def validate_url(self, url: str) -> bool:
        """
        Validate that URL is from Discord CDN for security.
        
        Args:
            url: Attachment URL to validate
            
        Returns:
            True if URL is valid and from Discord CDN
        """
        try:
            parsed = urlparse(url)
            return parsed.netloc in self.DISCORD_CDN_DOMAINS
        except Exception as e:
            logger.warning(f"Failed to parse attachment URL: {e}")
            return False
    
    def detect_file_type(self, filename: str, content_type: str) -> str:
        """
        Detect file type based on extension and content type.
        
        Args:
            filename: Name of the file
            content_type: MIME type
            
        Returns:
            File type: "text", "image", "pdf", "docx", "document", "other"
        """
        # Get extension
        ext = None
        if '.' in filename:
            ext = '.' + filename.rsplit('.', 1)[1].lower()
        
        # Check by extension first
        if ext in self.TEXT_EXTENSIONS:
            return "text"
        elif ext in self.IMAGE_EXTENSIONS:
            return "image"
        elif ext == '.pdf':
            return "pdf"
        elif ext in ['.docx', '.doc']:
            return "docx"
        elif ext in self.OTHER_DOCUMENT_EXTENSIONS:
            return "document"
        
        # Check by content type
        if content_type:
            content_type_lower = content_type.lower()
            if content_type_lower.startswith('text/'):
                return "text"
            elif content_type_lower.startswith('image/'):
                return "image"
            elif 'pdf' in content_type_lower:
                return "pdf"
            elif 'word' in content_type_lower or 'document' in content_type_lower:
                return "docx"
        
        return "other"
    
    async def download_file(
        self,
        url: str,
        max_size_mb: int,
        message_id: Optional[str] = None,
        filename: Optional[str] = None,
        context: Optional[Dict[str, Any]] = None
    ) -> Optional[bytes]:
        """
        Download file from URL with timeout and size validation.
        Supports automatic URL refresh on 404 if message_id is provided.
        
        Args:
            url: File URL
            max_size_mb: Maximum allowed size in MB
            message_id: Optional Discord message ID for URL refresh on 404
            filename: Optional filename for URL refresh
            context: Optional context for URL refresh (bot_client, channel_id)
            
        Returns:
            File data as bytes, or None if download failed
        """
        try:
            async with create_http_session(timeout_total=self.download_timeout) as session:
                async with session.get(url) as response:
                    if response.status == 404 and message_id and filename and context:
                        # URL expired, try to re-fetch fresh URL
                        logger.info(f"URL expired (404), attempting to re-fetch for {filename}")
                        from AI.tools.attachment_tools import _refetch_attachment_url
                        
                        fresh_url = await _refetch_attachment_url(message_id, filename, context)
                        if fresh_url:
                            logger.info(f"Got fresh URL, retrying download for {filename}")
                            # Retry with fresh URL
                            async with session.get(fresh_url) as retry_response:
                                if retry_response.status != 200:
                                    logger.warning(f"Failed to download file with fresh URL: HTTP {retry_response.status}")
                                    return None
                                
                                # Check content length
                                content_length = int(retry_response.headers.get('Content-Length', 0))
                                max_size_bytes = max_size_mb * 1024 * 1024
                                
                                if content_length > max_size_bytes:
                                    logger.warning(
                                        f"File too large: {content_length / (1024*1024):.1f}MB "
                                        f"(max: {max_size_mb}MB)"
                                    )
                                    return None
                                
                                # Download with fresh URL
                                data = await retry_response.read()
                                return data
                        else:
                            logger.warning(f"Failed to re-fetch URL for {filename}")
                            return None
                    
                    if response.status != 200:
                        logger.warning(f"Failed to download file: HTTP {response.status}")
                        return None
                    
                    # Check content length
                    content_length = int(response.headers.get('Content-Length', 0))
                    max_size_bytes = max_size_mb * 1024 * 1024
                    
                    if content_length > max_size_bytes:
                        logger.warning(
                            f"File too large: {content_length / (1024*1024):.1f}MB "
                            f"(max: {max_size_mb}MB)"
                        )
                        return None
                    
                    # Download
                    data = await response.read()
                    return data
                    
        except asyncio.TimeoutError:
            logger.warning(f"File download timeout after {self.download_timeout}s")
            return None
        except Exception as e:
            logger.error(f"Error downloading file: {e}")
            return None
    
    async def process_text_file(
        self,
        data: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """
        Process text file by decoding to UTF-8.
        
        Args:
            data: Raw file bytes
            filename: Name of the file
            
        Returns:
            Dict with content and metadata
        """
        try:
            # Try UTF-8 first
            try:
                content = data.decode('utf-8')
                encoding = 'utf-8'
            except UnicodeDecodeError:
                # Try other encodings
                for enc in ['latin-1', 'cp1252', 'iso-8859-1']:
                    try:
                        content = data.decode(enc)
                        encoding = enc
                        logger.debug(f"Decoded {filename} with {enc}")
                        break
                    except UnicodeDecodeError:
                        continue
                else:
                    return {
                        "error": "Failed to decode text file (unsupported encoding)",
                        "filename": filename
                    }
            
            return {
                "content": content,
                "encoding": encoding,
                "size": len(data),
                "lines": content.count('\n') + 1
            }
            
        except Exception as e:
            logger.error(f"Error processing text file {filename}: {e}")
            return {
                "error": f"Failed to process text file: {str(e)}",
                "filename": filename
            }
    
    async def process_image_file(
        self,
        attachment: Dict[str, Any]
    ) -> Dict[str, Any]:
        """
        Process image file using existing ImageProcessor.
        
        Args:
            attachment: Discord attachment dict
            
        Returns:
            Dict with base64 data and metadata
        """
        try:
            from utils.media_processor import ImageProcessor
            
            processor = ImageProcessor()
            
            # Create vision config
            config = {
                'max_image_size': self.max_image_size_mb,
                'vision_detail': 'auto'
            }
            
            result = await processor.process_image(attachment, config)
            
            if result:
                return {
                    "base64": result['base64'],
                    "format": result['format'],
                    "detail": result['detail'],
                    "size": result['size'],
                    "_vision_image": True  # Marker for vision processing
                }
            else:
                return {
                    "error": "Failed to process image",
                    "filename": attachment.get('filename', 'unknown')
                }
                
        except Exception as e:
            logger.error(f"Error processing image: {e}")
            return {
                "error": f"Failed to process image: {str(e)}",
                "filename": attachment.get('filename', 'unknown')
            }
    
    async def process_pdf_file(
        self,
        data: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """
        Process PDF file and extract text using PyPDF2.
        
        Args:
            data: Raw PDF bytes
            filename: Name of the file
            
        Returns:
            Dict with extracted text and metadata
        """
        try:
            import PyPDF2
            
            # Create PDF reader from bytes
            pdf_file = io.BytesIO(data)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract text from all pages
            text_parts = []
            num_pages = len(pdf_reader.pages)
            
            for page_num in range(num_pages):
                try:
                    page = pdf_reader.pages[page_num]
                    text = page.extract_text()
                    if text.strip():
                        text_parts.append(text)
                except Exception as e:
                    logger.warning(f"Failed to extract text from page {page_num + 1}: {e}")
                    continue
            
            extracted_text = '\n\n'.join(text_parts)
            
            # Get metadata
            metadata = {}
            if pdf_reader.metadata:
                try:
                    metadata = {
                        'title': pdf_reader.metadata.get('/Title', ''),
                        'author': pdf_reader.metadata.get('/Author', ''),
                        'subject': pdf_reader.metadata.get('/Subject', ''),
                        'creator': pdf_reader.metadata.get('/Creator', '')
                    }
                    # Remove empty values
                    metadata = {k: v for k, v in metadata.items() if v}
                except Exception:
                    pass
            
            return {
                "extracted_text": extracted_text,
                "pages": num_pages,
                "size": len(data),
                "metadata": metadata if metadata else None,
                "text_length": len(extracted_text)
            }
            
        except ImportError:
            logger.error("PyPDF2 not installed. Install with: pip install PyPDF2")
            return {
                "error": "PyPDF2 not installed",
                "filename": filename
            }
        except Exception as e:
            logger.error(f"Error processing PDF {filename}: {e}")
            return {
                "error": f"Failed to process PDF: {str(e)}",
                "filename": filename
            }
    
    async def process_docx_file(
        self,
        data: bytes,
        filename: str
    ) -> Dict[str, Any]:
        """
        Process DOCX file and extract text using python-docx.
        
        Args:
            data: Raw DOCX bytes
            filename: Name of the file
            
        Returns:
            Dict with extracted text and metadata
        """
        try:
            import docx
            
            # Create document from bytes
            docx_file = io.BytesIO(data)
            doc = docx.Document(docx_file)
            
            # Extract text from paragraphs
            paragraphs = []
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    paragraphs.append(text)
            
            # Extract text from tables
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = ' | '.join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        tables_text.append(row_text)
            
            # Combine all text
            extracted_text = '\n\n'.join(paragraphs)
            if tables_text:
                extracted_text += '\n\n--- Tables ---\n\n' + '\n'.join(tables_text)
            
            # Get metadata
            metadata = {}
            try:
                core_props = doc.core_properties
                metadata = {
                    'title': core_props.title or '',
                    'author': core_props.author or '',
                    'subject': core_props.subject or '',
                    'created': str(core_props.created) if core_props.created else '',
                    'modified': str(core_props.modified) if core_props.modified else ''
                }
                # Remove empty values
                metadata = {k: v for k, v in metadata.items() if v}
            except Exception:
                pass
            
            return {
                "extracted_text": extracted_text,
                "paragraphs": len(paragraphs),
                "tables": len(doc.tables),
                "size": len(data),
                "metadata": metadata if metadata else None,
                "text_length": len(extracted_text)
            }
            
        except ImportError:
            logger.error("python-docx not installed. Install with: pip install python-docx")
            return {
                "error": "python-docx not installed",
                "filename": filename
            }
        except Exception as e:
            logger.error(f"Error processing DOCX {filename}: {e}")
            return {
                "error": f"Failed to process DOCX: {str(e)}",
                "filename": filename
            }
    
    async def process_attachment(
        self,
        attachment: Dict[str, Any],
        include_content: bool = True,
        context: Optional[Dict[str, Any]] = None
    ) -> Dict[str, Any]:
        """
        Process a Discord attachment based on its type.
        
        Args:
            attachment: Discord attachment dict with url, filename, content_type, size, message_id
            include_content: Whether to include file content (default: True)
            context: Optional context for URL re-fetch (bot_client, channel_id)
            
        Returns:
            Dict with processed attachment data
        """
        url = attachment.get('url')
        filename = attachment.get('filename', 'unknown')
        content_type = attachment.get('content_type', '')
        size = attachment.get('size', 0)
        message_id = attachment.get('message_id')
        
        # Validate URL
        if not self.validate_url(url):
            return {
                "filename": filename,
                "error": "Invalid URL (not from Discord CDN)",
                "url": url
            }
        
        # Detect file type
        file_type = self.detect_file_type(filename, content_type)
        
        # Base result
        result = {
            "filename": filename,
            "content_type": content_type,
            "size": size,
            "url": url,
            "file_type": file_type
        }
        
        # If content not requested, return metadata only
        if not include_content:
            return result
        
        # Process based on file type
        if file_type == "text":
            # Check size limit
            max_size_bytes = self.max_text_size_kb * 1024
            if size > max_size_bytes:
                result["error"] = f"Text file too large: {size / 1024:.1f}KB (max: {self.max_text_size_kb}KB)"
                return result
            
            # Download and process
            data = await self.download_file(
                url,
                self.max_text_size_kb / 1024,
                message_id=message_id,
                filename=filename,
                context=context
            )
            if data:
                text_result = await self.process_text_file(data, filename)
                result.update(text_result)
            else:
                result["error"] = "Failed to download text file"
        
        elif file_type == "image":
            # Process image
            image_result = await self.process_image_file(attachment)
            result.update(image_result)
        
        elif file_type == "pdf":
            # Download and process PDF
            data = await self.download_file(
                url,
                self.max_document_size_mb,
                message_id=message_id,
                filename=filename,
                context=context
            )
            if data:
                pdf_result = await self.process_pdf_file(data, filename)
                result.update(pdf_result)
            else:
                result["error"] = "Failed to download PDF"
        
        elif file_type == "docx":
            # Download and process DOCX
            data = await self.download_file(
                url,
                self.max_document_size_mb,
                message_id=message_id,
                filename=filename,
                context=context
            )
            if data:
                docx_result = await self.process_docx_file(data, filename)
                result.update(docx_result)
            else:
                result["error"] = "Failed to download DOCX"
        
        elif file_type == "document":
            # Other documents - metadata only
            result["note"] = "Document type not supported for text extraction (metadata only)"
        
        else:
            # Other files - metadata only
            result["note"] = "File type not supported for content extraction (metadata only)"
        
        return result


# Global processor instance
_global_processor: Optional[AttachmentProcessor] = None


def get_attachment_processor() -> AttachmentProcessor:
    """Get the global attachment processor instance."""
    global _global_processor
    if _global_processor is None:
        _global_processor = AttachmentProcessor()
    return _global_processor
